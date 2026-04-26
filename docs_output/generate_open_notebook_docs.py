"""Generate Open-Notebook technical documentation (medium length) as a .docx file."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = Path(__file__).parent / "Open-Notebook_Documentation.docx"

doc = Document()
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

def h(text, level=1):
    return doc.add_heading(text, level=level)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    doc.add_paragraph(text, style="List Number")

def code(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, col in enumerate(headers):
        hdr[i].text = col
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
    doc.add_paragraph()

# --- Cover ----------------------------------------------------------------
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Open Notebook")
tr.bold = True
tr.font.size = Pt(30)
tr.font.color.rgb = RGBColor(0x0B, 0x3D, 0x91)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Technical Documentation")
sr.italic = True
sr.font.size = Pt(16)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("A privacy-focused, self-hosted alternative to Google Notebook LM\nApril 2026")
mr.italic = True

doc.add_paragraph()
para(
    "This document is a single technical reference for Open Notebook "
    "(lfnovo/open-notebook). It is aimed at developers, operators and "
    "stakeholders and covers architecture, modules, configuration, "
    "deployment, operation and extension."
)

doc.add_page_break()

# --- TOC ------------------------------------------------------------------
h("Table of contents", 1)
for line in [
    "1. Overview",
    "2. High-level architecture",
    "3. Backend API",
    "4. Core package (open_notebook/)",
    "5. AI providers and model selection",
    "6. Database (SurrealDB)",
    "7. Podcast generation",
    "8. Frontend",
    "9. Commands, workers and jobs",
    "10. Configuration reference",
    "11. Deployment (Docker Compose and single-container)",
    "12. Operating the system",
    "13. Extending Open Notebook",
    "14. Troubleshooting and operations",
    "15. Glossary",
]:
    bullet(line)

doc.add_page_break()

# 1. OVERVIEW --------------------------------------------------------------
h("1. Overview", 1)
para(
    "Open Notebook is an open-source, self-hosted alternative to Google "
    "Notebook LM. It lets a user collect heterogeneous sources (PDFs, web "
    "pages, audio, video, plain text), organise them in notebooks, chat "
    "with them through a Large Language Model, transform them, search "
    "them semantically, and generate multi-speaker podcasts from them. "
    "The project is designed around privacy and provider freedom: the "
    "user chooses which AI providers to use and all data stays on the "
    "machines the user controls."
)

h("1.1 Key facts", 2)
bullet("Repository: lfnovo/open-notebook (MIT).")
bullet("Current version: 1.8.1 (2026-03-10).")
bullet("Python: >=3.11,<3.13.")
bullet("Database: SurrealDB (graph + document + embedded vectors).")
bullet("Frontend: Next.js 16 / React 19 / TypeScript / Tailwind / Shadcn-UI.")
bullet("Backend: FastAPI (port 5055).")
bullet("Multi-provider AI abstraction: Esperanto + LangChain + LangGraph.")
bullet("Async job worker: surreal-commands-worker.")

h("1.2 Audience", 2)
bullet("Developers - extend routers, graphs, podcast flows, add providers.")
bullet("Operators - deploy, configure credentials, manage upgrades.")
bullet("Stakeholders - understand capabilities, data flow and comparison with Notebook LM.")

h("1.3 Positioning vs Google Notebook LM", 2)
table(
    ["Feature", "Open Notebook", "Google Notebook LM"],
    [
        ["Privacy", "Self-hosted, user owns data.", "Google cloud only."],
        ["AI providers", "16+ via Esperanto / LangChain.", "Google models only."],
        ["Podcast speakers", "1-4 with custom profiles.", "2 speakers, fixed format."],
        ["API access", "Full REST API (port 5055).", "No API."],
        ["Deployment", "Docker, single container, local or cloud.", "Hosted service."],
        ["Customisation", "Open source; transformations, prompts, graphs.", "Closed."],
        ["Cost model", "Pay per provider usage.", "Subscription + free tier."],
    ],
)

# 2. ARCHITECTURE ----------------------------------------------------------
h("2. High-level architecture", 1)
para(
    "Open Notebook is a three-tier application plus an asynchronous worker. "
    "All components communicate through well-defined boundaries so they "
    "can be scaled or replaced independently."
)

h("2.1 Tiers", 2)
table(
    ["Tier", "Responsibility", "Technology"],
    [
        ["Frontend", "UI, session, routing, API calls.", "Next.js 16 + React 19 on port 8502."],
        ["Backend API", "Business logic, LLM orchestration, persistence.", "FastAPI + Pydantic on port 5055."],
        ["Database", "Storage, vectors, relations, migrations.", "SurrealDB v2 on port 8000."],
        ["Worker", "Async jobs (embeddings, podcasts, long ingestion).", "surreal-commands-worker."],
    ],
)

h("2.2 Logical layers inside the backend", 2)
bullet("api/ - FastAPI app: routers, services, middleware, models.")
bullet("open_notebook/ai/ - LLM and embedding provider abstraction (Esperanto).")
bullet("open_notebook/database/ - SurrealDB client, repository helpers and migrations.")
bullet("open_notebook/domain/ - Pydantic entities bound to database tables.")
bullet("open_notebook/graphs/ - LangGraph state machines (chat, ask, source ingestion, transformations).")
bullet("open_notebook/podcasts/ - Podcast generation flows.")
bullet("open_notebook/search/ - Semantic search (OpenSearch integration).")
bullet("open_notebook/security/ - Encryption and authentication helpers.")

h("2.3 Process topology", 2)
bullet("Typical stack: SurrealDB, API, worker and frontend - four processes.")
bullet("docker-compose.yml runs SurrealDB and a combined Open Notebook image that runs API + worker + frontend via supervisord.")
bullet("Dockerfile.single packages everything including SurrealDB into a single image for minimal deployments.")

# 3. BACKEND API -----------------------------------------------------------
h("3. Backend API", 1)
para(
    "The API is defined in api/ and started with run_api.py (which boots "
    "uvicorn on api.main:app). Default bind is 127.0.0.1:5055; host and "
    "port are overridable through API_HOST and API_PORT."
)

h("3.1 Routers", 2)
table(
    ["Router", "Main endpoints", "Purpose"],
    [
        ["notebooks.py", "POST / GET / DELETE /notebooks", "Notebook CRUD."],
        ["sources.py", "POST / GET / DELETE /sources/{id}", "Ingest and manage files, URLs, audio and video."],
        ["notes.py", "POST / GET /notes/{id}", "User notes tied to sources."],
        ["chat.py", "POST /chat", "Global chat with notebook context."],
        ["source_chat.py", "POST /sources/{source_id}/chat", "Chat scoped to a specific source."],
        ["podcasts.py", "POST / GET /podcasts/{id}, /episodes/{id}/retry", "Outline, transcript and audio generation."],
        ["models.py", "GET /models, POST /models/config", "AI model registry and provisioning."],
        ["credentials.py", "CRUD, /test, /migrate", "Encrypted storage of provider API keys."],
        ["transformations.py", "POST /transformations", "Apply user-defined transformations."],
        ["insights.py", "GET /sources/{id}/insights", "Generated insights per source."],
        ["search.py", "POST /search", "Semantic search over sources and notes."],
        ["auth.py", "POST /auth/password", "Password-based authentication."],
        ["languages.py", "GET /languages", "Available podcast languages."],
        ["commands.py", "GET /commands/{command_id}", "Status of async jobs."],
        ["vision.py", "POST /vision/image-analysis, /vision/video-tracking", "Image analysis and video object tracking. Form field 'engine' selects 'sam3' (open-vocabulary; query/target required) or 'rfdetr' (closed-vocabulary COCO; query/target optional). Forwards uploads to NOVA-Researcher."],
    ],
)

h("3.2 Services layer", 2)
bullet("chat_service.py - invokes LangGraph chat graph and persists message history (SQLite checkpointer).")
bullet("podcast_service.py - orchestrates outline + transcript + TTS.")
bullet("sources_service.py - ingest, chunk, embed and persist sources.")
bullet("models_service.py - manages provider / model configuration and defaults.")
bullet("notes_service.py, insights_service.py, transformations_service.py, episode_profiles_service.py, etc.")

h("3.3 Cross-cutting concerns", 2)
bullet("Async everywhere: FastAPI, SurrealDB client, LangGraph graphs.")
bullet("Authentication: optional password middleware (OPEN_NOTEBOOK_PASSWORD).")
bullet("Exception mapping to HTTP codes via custom handlers.")
bullet("Schema migrations run on startup via the lifespan handler.")

# 4. CORE PACKAGE ----------------------------------------------------------
h("4. Core package (open_notebook/)", 1)
table(
    ["Module", "Key files", "Purpose"],
    [
        ["ai/", "provision.py, models.py, key_provider.py, connection_tester.py, vision.py, model_discovery.py", "Provider-agnostic LLM/embedding/TTS/STT wrappers."],
        ["database/", "db.py, repository.py, migrate.py, async_migrate.py, migrations/", "SurrealDB driver wrappers and cumulative SurrealQL migrations."],
        ["domain/", "notebook.py, source.py, note.py, credential.py, transformation.py, audit_log.py, provider_config.py, role.py, content_settings.py, user.py", "Pydantic entities bound to SurrealDB tables."],
        ["graphs/", "chat.py, ask.py, source.py, source_chat.py, transformation.py, tools.py, prompt.py", "LangGraph state machines used by services."],
        ["podcasts/", "models.py, migration.py", "Podcast and speaker profile entities."],
        ["search/", "-", "OpenSearch-based vector search integration."],
        ["security/", "-", "Encryption of stored credentials (SHA-256 keyed with OPEN_NOTEBOOK_ENCRYPTION_KEY)."],
        ["utils/", "-", "Text, validation, i18n helpers (pycountry, babel)."],
        ["config.py", "-", "Environment-driven runtime settings."],
    ],
)

# 5. AI PROVIDERS ----------------------------------------------------------
h("5. AI providers and model selection", 1)
para(
    "Model access is unified through the Esperanto library, with LangChain "
    "adapters per provider. Users can mix and match providers - for example, "
    "use Anthropic for chat, OpenAI for embeddings, Ollama for a local "
    "fallback and ElevenLabs for TTS."
)

h("5.1 Supported providers", 2)
bullet("OpenAI and OpenAI-compatible servers.")
bullet("Anthropic (Claude).")
bullet("Google Generative AI (Gemini).")
bullet("Groq.")
bullet("Mistral.")
bullet("DeepSeek.")
bullet("Hugging Face.")
bullet("Ollama - local open-weights models.")
bullet("LM Studio and Speaches (via OpenAI-compatible interface).")

h("5.2 Model roles", 2)
bullet("chat / instruct - conversations and transformations.")
bullet("embedding - vector embeddings for sources and notes.")
bullet("TTS - podcast voice synthesis.")
bullet("STT - speech-to-text for audio/video sources.")
bullet("large_context_model - used automatically when the context exceeds ~105 000 tokens.")

h("5.3 Selection logic (provision_langchain_model)", 2)
numbered("If the request carries an explicit model_id, use it.")
numbered("If the estimated context exceeds the large-context threshold, use the configured large_context_model.")
numbered("Otherwise fall back to the default model registered for the requested role.")

h("5.4 Credential management", 2)
bullet("API keys are stored in the credential table.")
bullet("Values are encrypted at rest using OPEN_NOTEBOOK_ENCRYPTION_KEY (SHA-256 derived key).")
bullet("Credentials can be tested from the UI or via /credentials/test before being used.")

# 6. DATABASE --------------------------------------------------------------
h("6. Database (SurrealDB)", 1)
para(
    "SurrealDB provides document storage, graph relations and vector "
    "similarity in a single engine. Open Notebook uses its native features "
    "for embeddings so that sources, chunks and notes can be retrieved by "
    "content similarity without an additional vector database."
)

h("6.1 Key tables (inferred from domain)", 2)
bullet("notebook - user-owned collection of sources and notes.")
bullet("source - ingested content (PDF, URL, audio, video, text).")
bullet("note - user or AI-generated note attached to a source.")
bullet("chat_session - chat threads.")
bullet("credential - encrypted provider API keys.")
bullet("model - LLM / embedding / TTS / STT registry.")
bullet("transformation - user-defined content transformations.")
bullet("episode_profile, speaker_profile - podcast templates and voices.")
bullet("audit_log - user audit trail.")
bullet("reference - graph edges between notebook/source/note.")

h("6.2 Migrations", 2)
bullet("19 SurrealQL migration files (migrations/1.surrealql to 19.surrealql) with matching *_down.surrealql files.")
bullet("AsyncMigrationManager runs pending migrations at API start-up.")
bullet("Schema upgrades are additive and reversible via down scripts.")

h("6.3 Connection", 2)
bullet("Driver: surrealdb (Python, async).")
bullet("Transport: WebSocket (ws:// or wss://).")
bullet("Configured via SURREAL_URL, SURREAL_USER, SURREAL_PASS (or SURREAL_PASSWORD), SURREAL_NAMESPACE, SURREAL_DATABASE.")

# 7. PODCAST ---------------------------------------------------------------
h("7. Podcast generation", 1)
para(
    "Podcast generation is delegated to the podcast-creator library "
    "(>=0.12.0) and composed inside open_notebook/podcasts/. A generation "
    "job is an asynchronous flow broken into three phases."
)
numbered("Outline - the LLM produces a multi-speaker outline from one or more sources using prompts/podcast/outline.jinja.")
numbered("Transcript - the outline is expanded into a full dialogue using prompts/podcast/transcript.jinja.")
numbered("Audio synthesis - the transcript is rendered to speech per speaker using the chosen TTS provider.")

h("7.1 TTS providers", 2)
bullet("OpenAI TTS (OpenAI API).")
bullet("ElevenLabs (premium, high-quality voices).")
bullet("Google Cloud TTS.")
bullet("Speaches or any OpenAI-compatible TTS server (local, free).")

h("7.2 Profiles", 2)
bullet("SpeakerProfile - voice id, language, provider, voice style.")
bullet("EpisodeProfile - speaker list, outline model, transcript model, defaults.")
bullet("Profiles are stored in SurrealDB and editable from the UI.")

h("7.3 Execution", 2)
bullet("Jobs are dispatched to the async worker (surreal-commands-worker).")
bullet("Status is exposed through /commands/{command_id}.")
bullet("Audio output is persisted on the data volume (/app/data).")

# 8. FRONTEND --------------------------------------------------------------
h("8. Frontend", 1)
para(
    "The frontend lives under frontend/ and is a Next.js 16 application "
    "using the App Router, React 19, TypeScript, Tailwind CSS and "
    "Shadcn-UI components. It listens on port 8502 in production "
    "(supervisord) and port 3000 in local dev."
)

h("8.1 Main routes", 2)
bullet("/notebooks - list and detail.")
bullet("/sources - source management and ingestion.")
bullet("/notes - note editor.")
bullet("/chat - global chat with notebook context.")
bullet("/research - research tooling (can integrate with an external researcher, e.g. NOVA-Researcher via NOVA_RESEARCHER_URL).")
bullet("/search - semantic search.")
bullet("/podcasts - podcast profiles and generation.")
bullet("/transformations - create and run transformations.")
bullet("/vision - image analysis and video tracking, with two engines: SAM 3 (open-vocabulary, requires a text prompt) and RF-DETR (closed-vocabulary COCO, prompt-free; an optional class name is used as a filter).")
bullet("/settings, /advanced, /admin - configuration and administration.")

h("8.2 State and data", 2)
bullet("State: Zustand.")
bullet("Data fetching: TanStack Query with a thin API client in lib/.")
bullet("Internationalisation: babel + pycountry server-side; multiple UI languages (English, Portuguese, Chinese Simplified & Traditional, Japanese, Russian, Bengali).")
bullet("Tests: Vitest.")

# 9. COMMANDS --------------------------------------------------------------
h("9. Commands, workers and jobs", 1)

h("9.1 Asynchronous worker", 2)
bullet("Implementation: surreal-commands-worker (external library, backed by SurrealDB).")
bullet("Consumes jobs enqueued by API services (podcast generation, embedding rebuild, long ingestion).")
bullet("Runs alongside the API under supervisord.")

h("9.2 CLI commands (commands/)", 2)
bullet("embedding_commands.py - (re)build embeddings and manage vector indices.")
bullet("podcast_commands.py - generate outlines, transcripts and episodes from the CLI.")
bullet("source_commands.py - import / export sources in bulk.")
bullet("opensearch_commands.py - maintain OpenSearch indices.")

h("9.3 Makefile targets", 2)
table(
    ["Target", "Purpose"],
    [
        ["database", "Start SurrealDB via docker compose."],
        ["run / frontend", "npm run dev for the frontend."],
        ["lint", "mypy type-checking."],
        ["ruff", "ruff check . --fix."],
        ["start-all", "Bring up API + worker + frontend."],
        ["worker", "Start the async worker."],
        ["docker-build-local", "Build the Docker image locally."],
        ["docker-release", "Multi-platform build and push (Docker Hub + GHCR)."],
        ["docker-push", "Push already-built image."],
    ],
)

# 10. CONFIGURATION --------------------------------------------------------
h("10. Configuration reference", 1)
para(
    "Configuration is driven by environment variables. The docker-compose "
    "file supplies sensible defaults; only OPEN_NOTEBOOK_ENCRYPTION_KEY is "
    "strictly required."
)

h("10.1 Database", 2)
table(
    ["Variable", "Example", "Purpose"],
    [
        ["SURREAL_URL", "ws://surrealdb:8000/rpc", "SurrealDB WebSocket endpoint."],
        ["SURREAL_USER", "root", "SurrealDB user."],
        ["SURREAL_PASSWORD (or SURREAL_PASS)", "root", "SurrealDB password."],
        ["SURREAL_NAMESPACE", "open_notebook", "Namespace name."],
        ["SURREAL_DATABASE", "open_notebook", "Database name."],
    ],
)

h("10.2 API server", 2)
table(
    ["Variable", "Default", "Purpose"],
    [
        ["API_HOST", "127.0.0.1", "Bind address."],
        ["API_PORT", "5055", "Listen port."],
        ["API_BASE_URL", "http://127.0.0.1:5055", "Absolute URL used by the frontend."],
        ["API_RELOAD", "true (dev)", "Uvicorn auto-reload."],
        ["API_CLIENT_TIMEOUT", "300.0", "HTTP client timeout in seconds."],
    ],
)

h("10.3 Security", 2)
table(
    ["Variable", "Required", "Purpose"],
    [
        ["OPEN_NOTEBOOK_ENCRYPTION_KEY", "Yes", "Key used to encrypt stored credentials."],
        ["OPEN_NOTEBOOK_PASSWORD", "No", "Password for UI authentication; empty disables auth."],
    ],
)

h("10.4 Optional integrations", 2)
table(
    ["Variable", "Purpose"],
    [
        ["NEXT_PUBLIC_API_URL", "API URL seen by the browser."],
        ["INTERNAL_API_URL", "API URL used by Next.js server-side (ISR, actions)."],
        ["NOVA_RESEARCHER_URL", "External researcher endpoint (e.g. http://localhost:8002)."],
    ],
)

# 11. DEPLOYMENT -----------------------------------------------------------
h("11. Deployment", 1)

h("11.1 Docker Compose", 2)
para("Recommended deployment. The compose file ships two services:")
table(
    ["Service", "Image", "Ports", "Purpose"],
    [
        ["surrealdb", "surrealdb/surrealdb:v2", "8000", "Database."],
        ["open_notebook", "lfnovo/open_notebook:v1-latest", "5055 (API), 8502 (UI)", "API + worker + frontend under supervisord."],
    ],
)
para("Minimum configuration:")
code(
    "services:\n"
    "  surrealdb:\n"
    "    image: surrealdb/surrealdb:v2\n"
    "    command: start --log info --user root --pass root rocksdb:/mydata/mydatabase.db\n"
    "    ports: [\"8000:8000\"]\n"
    "    volumes: [\"./surreal_data:/mydata\"]\n"
    "    restart: always\n"
    "\n"
    "  open_notebook:\n"
    "    image: lfnovo/open_notebook:v1-latest\n"
    "    ports: [\"8502:8502\", \"5055:5055\"]\n"
    "    environment:\n"
    "      - OPEN_NOTEBOOK_ENCRYPTION_KEY=change-me\n"
    "      - SURREAL_URL=ws://surrealdb:8000/rpc\n"
    "      - SURREAL_USER=root\n"
    "      - SURREAL_PASSWORD=root\n"
    "      - SURREAL_NAMESPACE=open_notebook\n"
    "      - SURREAL_DATABASE=open_notebook\n"
    "    volumes: [\"./notebook_data:/app/data\"]\n"
    "    depends_on: [surrealdb]\n"
    "    restart: always\n"
)

h("11.2 Single-container image (Dockerfile.single)", 2)
bullet("Bundles SurrealDB, API, worker and frontend in one image.")
bullet("Managed by supervisord (surrealdb, api, worker, frontend) with explicit priorities 5 / 10 / 20 / 30.")
bullet("Exposes 5055 (API), 8502 (UI) and 8000 (SurrealDB internal).")
bullet("Ideal for small single-host deployments where one container is preferred over compose.")

h("11.3 Build image (standard Dockerfile)", 2)
bullet("Multi-stage build on python:3.12-slim.")
bullet("Installs uv and Node.js, builds frontend and backend.")
bullet("Pre-downloads tiktoken encoding into /app/tiktoken-cache (offline resilience).")
bullet("Installs ffmpeg and supervisor at runtime.")
bullet("Multi-arch builds (linux/amd64, linux/arm64) via buildx (make docker-release).")

h("11.4 Process supervision", 2)
para(
    "Both supervisord.conf (compose) and supervisord.single.conf (single-container) "
    "define four processes with auto-restart. api starts first, worker waits "
    "briefly, and frontend waits for the API through a wait-for-api.sh "
    "script. In single-container mode, surrealdb starts first (priority 5)."
)

# 12. OPERATING ------------------------------------------------------------
h("12. Operating the system", 1)

h("12.1 First start", 2)
numbered("docker compose up -d.")
numbered("Open http://localhost:8502.")
numbered("If OPEN_NOTEBOOK_PASSWORD is set, log in with that password.")
numbered("Go to Settings / Models, add API keys for the AI providers you want to use; keys are encrypted with OPEN_NOTEBOOK_ENCRYPTION_KEY.")
numbered("Configure default chat / embedding / TTS / STT models.")

h("12.2 Typical workflow", 2)
numbered("Create a notebook.")
numbered("Add sources (PDF, URL, audio, video, text). Content extraction runs via content-core; transcripts via the STT provider.")
numbered("Sources are chunked, embedded and stored in SurrealDB.")
numbered("Ask questions in the chat; the LangGraph chat graph retrieves relevant chunks and composes a grounded answer.")
numbered("Apply transformations (summaries, extractions) or generate insights per source.")
numbered("Optionally generate a podcast - pick an Episode Profile, trigger generation, wait for the worker to produce outline, transcript and audio.")

h("12.3 Using the REST API", 2)
code(
    "# Create a notebook\n"
    "curl -X POST http://localhost:5055/notebooks \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"name\": \"Marinha Research\"}'\n"
    "\n"
    "# Upload a source (URL)\n"
    "curl -X POST http://localhost:5055/sources \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"notebook_id\": \"notebook:xxx\", \"type\": \"url\", \"url\": \"https://...\"}'\n"
    "\n"
    "# Chat with the notebook\n"
    "curl -X POST http://localhost:5055/chat \\\n"
    "  -H 'Content-Type: application/json' \\\n"
    "  -d '{\"notebook_id\": \"notebook:xxx\", \"message\": \"Summarise the latest source\"}'\n"
)

# 13. EXTENDING ------------------------------------------------------------
h("13. Extending Open Notebook", 1)

h("13.1 Adding a new API endpoint", 2)
numbered("Create a module under api/routers/ and define an APIRouter.")
numbered("Place business logic in a matching *_service.py under api/.")
numbered("Register the router in api/main.py.")
numbered("Add a service test under tests/ using pytest-asyncio.")

h("13.2 Adding a new AI provider", 2)
bullet("Prefer providers already supported by Esperanto and LangChain.")
bullet("Register the provider in open_notebook/ai/provision.py.")
bullet("Extend the model registry so the provider can be picked from the UI.")
bullet("Add credential tests to /credentials/test.")

h("13.3 Adding a new LangGraph workflow", 2)
bullet("Add a new graph module under open_notebook/graphs/.")
bullet("Expose it through a service and a router endpoint.")
bullet("Persist checkpoints with the existing SQLite checkpointer if the graph is conversational.")

h("13.4 New database entities", 2)
bullet("Add a new numbered SurrealQL migration (N.surrealql) plus its down file.")
bullet("Add a Pydantic model under open_notebook/domain/.")
bullet("Extend the repository helpers in open_notebook/database/repository.py.")

# 14. TROUBLESHOOTING ------------------------------------------------------
h("14. Troubleshooting and operations", 1)
table(
    ["Symptom", "Likely cause", "Action"],
    [
        ["UI cannot reach API", "API not up or NEXT_PUBLIC_API_URL wrong.", "Check supervisord logs; confirm /health on port 5055."],
        ["SurrealDB hangs at startup", "Volume or version mismatch.", "Clear ./surreal_data (dev only) or pin the surrealdb image to v2."],
        ["Credential test fails", "Wrong API key or expired plan.", "Re-enter key; check provider status."],
        ["Podcast job stuck 'running'", "Worker not running or TTS provider down.", "Check worker logs; retry via /episodes/{id}/retry."],
        ["Tiktoken network errors", "Offline container.", "Image pre-downloads encoding into /app/tiktoken-cache - verify it was not overwritten by a volume."],
        ["Wrong model used for long context", "large_context_model not configured.", "Set it in Settings / Models."],
        ["Decrypted credentials break", "OPEN_NOTEBOOK_ENCRYPTION_KEY changed.", "Restore original key or re-enter credentials."],
    ],
)

para("Logs:")
bullet("Docker: docker compose logs -f open_notebook.")
bullet("Inside the container: supervisord exposes api, worker and frontend logs on stdout.")
bullet("Python logging uses loguru and writes to stdout by default.")

# 15. GLOSSARY -------------------------------------------------------------
h("15. Glossary", 1)
table(
    ["Term", "Meaning"],
    [
        ["Notebook", "A user-owned collection of sources, notes and chats."],
        ["Source", "An ingested artefact (PDF, URL, audio, video, text)."],
        ["Note", "Free-form or AI-generated text attached to a source or notebook."],
        ["Transformation", "A reusable LLM operation that rewrites or extracts from content."],
        ["Insight", "An auto-generated analytical summary of a source."],
        ["Episode Profile", "Template for a podcast episode - speakers, models, language."],
        ["Speaker Profile", "TTS voice configuration."],
        ["Esperanto", "Python library providing a unified API across many AI providers."],
        ["LangGraph", "Framework for building stateful LLM workflows as graphs."],
        ["SurrealDB", "Multi-model database used for all persistent state and vectors."],
        ["Supervisord", "Process manager that keeps API, worker and frontend running."],
    ],
)

doc.save(OUT)
print(f"WROTE: {OUT}")
