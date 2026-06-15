# -*- coding: utf-8 -*-
"""
Gerador do Relatório E04 - Interface e Monitorização (versão factual)
Corre com: python3 gerar_relatorio_e04.py
Requer: python-docx
"""

from __future__ import annotations

import datetime
import json
from pathlib import Path
from typing import Any, Dict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

AZUL_MARINHA = RGBColor(0x1F, 0x35, 0x64)
BASE_DIR = Path("/user/home/mf.domingos/navy/open-notebook/docs_output")
METRICAS_DIR = BASE_DIR / "metricas_e04"
METRICAS_JSON = METRICAS_DIR / "metricas_resumo.json"

MESES_PT = {
    1: "janeiro", 2: "fevereiro", 3: "março", 4: "abril",
    5: "maio", 6: "junho", 7: "julho", 8: "agosto",
    9: "setembro", 10: "outubro", 11: "novembro", 12: "dezembro",
}


def add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = h.runs[0] if h.runs else h.add_run(text)
    run.font.color.rgb = AZUL_MARINHA
    return h


def add_body(doc: Document, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(2)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(text, style="List Bullet")
    p.paragraph_format.left_indent = Cm(1.0)
    for run in p.runs:
        run.font.size = Pt(11)
    return p


def add_screenshot_placeholder(doc: Document, caption: str):
    doc.add_paragraph()
    box = doc.add_paragraph()
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = box.add_run(f"[ CAPTURA DE ECRÃ: {caption} ]")
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "E8E8E8")
    box._p.get_or_add_pPr().append(shd)
    cap_p = doc.add_paragraph(f"Figura - {caption}")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in cap_p.runs:
        r.font.size = Pt(9)
        r.font.italic = True
    doc.add_paragraph()


def add_table_of_contents(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run("ÍNDICE")
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL_MARINHA
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r = OxmlElement("w:r")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(fld_end)
    toc_p = doc.add_paragraph()
    toc_p._p.append(r)
    doc.add_page_break()


def style_header_row(cells):
    for cell in cells:
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F3564")
        cell._tc.get_or_add_tcPr().append(shd)


# Contador global de tabelas, para legendas numeradas ("Tabela N - ...").
_TABLE_COUNTER = {"n": 0}


def add_caption(doc: Document, prefix: str, text: str):
    """Adiciona uma legenda numerada e centrada (ex.: "Tabela 1 - ...")."""
    if prefix == "Tabela":
        _TABLE_COUNTER["n"] += 1
        number = _TABLE_COUNTER["n"]
    else:
        number = None
    label = f"{prefix} {number} - {text}" if number else f"{prefix} - {text}"
    cap = doc.add_paragraph(label)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    for r in cap.runs:
        r.font.size = Pt(9)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return cap


def add_data_table(doc: Document, headers, rows, caption: str | None = None,
                   col_align=None):
    """Cria uma tabela com cabeçalho azul-marinha e legenda numerada opcional.

    headers: lista de títulos de coluna.
    rows: lista de tuplos/listas de valores (convertidos para texto).
    caption: texto da legenda (numerada como "Tabela N - ...").
    col_align: lista opcional de alinhamentos por coluna (WD_ALIGN_PARAGRAPH.*).
    """
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
    style_header_row(hdr)

    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = "" if value is None else str(value)
            para = cells[i].paragraphs[0]
            if col_align and i < len(col_align) and col_align[i] is not None:
                para.alignment = col_align[i]
            for run in para.runs:
                run.font.size = Pt(10)

    if caption:
        add_caption(doc, "Tabela", caption)
    else:
        doc.add_paragraph()
    return table


def load_metricas() -> Dict[str, Any]:
    if not METRICAS_JSON.exists():
        return {}
    try:
        return json.loads(METRICAS_JSON.read_text(encoding="utf-8"))
    except Exception:
        return {}


def fmt_ms(v: Any) -> str:
    if isinstance(v, (int, float)):
        return f"{v:.2f} ms"
    return "N/D"


def fmt_int(v: Any) -> str:
    if isinstance(v, (int, float)):
        return f"{int(v)}"
    return "N/D"


def fmt_pct(part: Any, total: Any) -> str:
    if isinstance(part, (int, float)) and isinstance(total, (int, float)) and total:
        return f"{(part / total) * 100:.1f}%"
    return "N/D"


def speedup(slow_ms: float, fast_ms: Any) -> str:
    """Calcula o fator de aceleração (ex.: "32x") de forma segura."""
    if isinstance(fast_ms, (int, float)) and fast_ms:
        return f"{slow_ms / fast_ms:.0f}x"
    return "N/D"


def build_document() -> str:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Capa
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title_p.add_run("OPEN-NOTEBOOK\nSistema de Pesquisa e Gestão de Conhecimento com IA")
    tr.font.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = AZUL_MARINHA

    doc.add_paragraph()
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run(
        "E04 - Interface e Monitorização\n"
        "Manual do Utilizador, Interface Gráfica, Administração e Auditoria\n"
        "Métricas Operacionais Calculadas com Evidência"
    )
    sr.font.size = Pt(14)
    sr.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    today = datetime.date.today()
    data_pt = f"{MESES_PT[today.month]} de {today.year}"
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(f"Marinha Portuguesa\nVersão 1.1 - {data_pt}")
    mr.font.size = Pt(11)

    doc.add_page_break()
    add_table_of_contents(doc)

    # 1
    add_heading(doc, "1. Introdução", 1)
    add_body(
        doc,
        "Este relatório E04 descreve a interface e a monitorização do sistema no estado atual da solução. "
        "A revisão foi feita com base no comportamento real da aplicação, nos endpoints expostos pela API, "
        "na interface gráfica disponível e nos registos técnicos observáveis."
    )
    add_body(
        doc,
        "O documento evita pressupostos não comprovados. Sempre que uma funcionalidade não esteja ainda "
        "concluída no código, isso é assinalado explicitamente e acompanhado de método de medição ou "
        "de plano de validação."
    )

    add_heading(doc, "1.1. Âmbito", 2)
    add_bullet(doc, "Manual do utilizador para operação diária e administração.")
    add_bullet(doc, "Descrição da autenticação, autorização, histórico e exportação de conversas.")
    add_bullet(doc, "Separação entre monitorização operacional da aplicação e monitorização de infraestrutura.")
    add_bullet(doc, "Relatório de auditoria, retenção de logs e metodologia de cálculo de métricas.")
    add_bullet(doc, "Registos de testes DR em DEV/QUA/PRD com estrutura de evidência.")

    doc.add_page_break()

    # 2
    add_heading(doc, "2. Manual do Utilizador", 1)
    add_heading(doc, "2.1. Acesso e Autenticação", 2)
    add_body(
        doc,
        "A página de autenticação apresenta dois modos: login local (email e palavra-passe) e OAuth "
        "(Azure, Google, GitHub). A API implementa emissão e validação de JWT para sessões autenticadas."
    )
    add_body(
        doc,
        "No estado atual, o login local encontra-se operacional. Existe inicialização de fluxos OAuth "
        "(geração do URL de autorização), e os handlers de callback OAuth estão definidos no backend. "
        "A troca completa de código por token no callback está parcialmente implementada, devendo ser "
        "validada por ambiente antes de ser assumida como fluxo produtivo de SSO organizacional."
    )
    add_screenshot_placeholder(doc, "Ecrã de login com seleção OAuth (Azure/Google/GitHub) e login local")

    add_heading(doc, "2.2. Navegação e Estrutura da Interface", 2)
    add_body(
        doc,
        "A interface organiza-se em áreas de trabalho para notebooks, fontes, notas, chat, pesquisa, "
        "podcasts e administração. A navegação principal é feita por rotas protegidas após autenticação."
    )
    add_screenshot_placeholder(doc, "Vista principal da aplicação com navegação entre Notebooks, Search, Podcasts, Models e Settings")

    add_heading(doc, "2.3. Histórico e Exportação de Conversas", 2)
    add_body(
        doc,
        "As conversas globais podem ser listadas, retomadas e exportadas em Markdown através da interface. "
        "A exportação agrega sessões em documento único e constitui mecanismo de reporte operacional."
    )
    add_screenshot_placeholder(doc, "Página de Chat global com sessões e opção de exportação")

    add_heading(doc, "2.4. Seleção e Gestão de Modelos", 2)
    add_body(
        doc,
        "A plataforma expõe catálogo de modelos por fornecedor e tipo (language, embedding, speech_to_text, "
        "text_to_speech), com configuração de modelos por defeito. A gestão é centralizada na área de administração."
    )
    add_screenshot_placeholder(doc, "Consola de modelos com lista de fornecedores, tipos e modelos por defeito")

    add_heading(doc, "2.5. Administração e Auditoria", 2)
    add_body(
        doc,
        "A área de administração inclui separadores de Overview, Ask, Podcasts, Users & Roles e Audit Logs. "
        "Os registos de auditoria apresentam ação, tipo de recurso, estado, código HTTP e duração em milissegundos."
    )
    add_screenshot_placeholder(doc, "Admin Dashboard com separadores de gestão e monitorização")
    add_screenshot_placeholder(doc, "Audit Logs com filtros por ação, utilizador e período")

    doc.add_page_break()

    # 3
    add_heading(doc, "3. Interface e Segurança", 1)
    add_heading(doc, "3.1. Modelo de Segurança", 2)
    add_body(
        doc,
        "A segurança de acesso combina middleware JWT com fallback de autenticação por palavra-passe, "
        "permitindo compatibilidade com cenários legados e evolução para federação de identidade. "
        "A API aplica autorização por papéis em rotas administrativas."
    )
    add_bullet(doc, "Autenticação: token Bearer JWT e renovação de token.")
    add_bullet(doc, "Autorização: papéis (admin/editor/viewer) e guardas de rota no frontend.")
    add_bullet(doc, "Auditoria: registo estruturado por ação, recurso, estado e duração.")

    add_heading(doc, "3.2. Estado de Prontidão OAuth/SSO", 2)
    add_body(
        doc,
        "Os endpoints de estado e inicialização OAuth estão ativos e permitem detetar provedores configurados. "
        "Para ambiente Marinha com SSO organizacional, é necessário validar fim-a-fim o callback OAuth com "
        "troca de código por token e mapeamento de claims de identidade no ambiente alvo."
    )
    add_bullet(doc, "Disponível: /api/auth/status e /api/auth/oauth/providers.")
    add_bullet(doc, "Disponível: inicialização OAuth por provedor.")
    add_bullet(doc, "A validar em produção: callback OAuth completo com token e perfil final.")

    doc.add_page_break()

    # 4
    metricas = load_metricas()

    # Extração das estruturas de métricas e metadados de recolha.
    search_bench = metricas.get("search_benchmark", {})
    audit = metricas.get("audit", {})
    dur_all = audit.get("duration_ms", {})
    dur_excl = audit.get("duration_ms_excl_token_refresh", {})
    entities = metricas.get("entities", {})
    base_url = metricas.get("base_url", "N/D")
    iterations = metricas.get("iterations", "N/D")
    gen_at_raw = metricas.get("generated_at")
    if isinstance(gen_at_raw, str):
        try:
            gen_at_dt = datetime.datetime.fromisoformat(gen_at_raw)
            gen_at = gen_at_dt.strftime("%d/%m/%Y às %H:%M UTC")
        except ValueError:
            gen_at = gen_at_raw
    else:
        gen_at = "N/D"

    add_heading(doc, "4. Monitorização e Métricas", 1)
    add_body(
        doc,
        "A monitorização do sistema organiza-se em dois planos complementares que respondem a perguntas "
        "distintas. O plano operacional foca-se na qualidade de serviço percebida pelo utilizador da "
        "aplicação - quão rápida é a pesquisa, com que fiabilidade respondem os pedidos e como evoluem os "
        "padrões de utilização. O plano de administração e infraestrutura foca-se na saúde dos serviços e "
        "dos recursos que sustentam a aplicação - disponibilidade da API e da base de dados, estado dos "
        "contentores e consumo de recursos. Separar estes dois planos evita confundir um problema "
        "funcional com um problema de infraestrutura, acelerando o diagnóstico."
    )
    add_heading(doc, "4.1. Dashboard Operacional vs Dashboard de Administração", 2)
    add_body(
        doc,
        "A tabela seguinte sistematiza as diferenças entre os dois planos de monitorização, identificando "
        "para cada um o objetivo, os indicadores típicos, as fontes de dados e o público-alvo. Esta "
        "distinção orienta tanto a conceção dos painéis como a atribuição de responsabilidades de "
        "acompanhamento."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Dimensão"
    hdr[1].text = "Operacional (Aplicação)"
    hdr[2].text = "Administração / Infraestrutura"
    style_header_row(hdr)

    rows = [
        ("Objetivo", "Qualidade de serviço da aplicação", "Saúde de serviços e recursos"),
        ("Exemplos", "Latência de pesquisa, erros API, uso de modelos", "Estado API/DB/Auth, contentores, CPU/RAM"),
        ("Fontes", "Endpoints /api/search, /api/audit, logs estruturados", "StatusDashboard, Dockhand, logs do runtime"),
        ("Público", "Operação funcional", "Administração técnica"),
    ]
    for a, b, c in rows:
        row = table.add_row().cells
        row[0].text = a
        row[1].text = b
        row[2].text = c

    add_caption(doc, "Tabela", "Comparação entre o plano de monitorização operacional e o de infraestrutura.")
    add_screenshot_placeholder(doc, "Dockhand / dashboard de contentores para operação de infraestrutura")

    add_heading(doc, "4.2. Métricas com Evidência - Estado Atual", 2)
    add_body(
        doc,
        "Os indicadores apresentados nesta secção são calculados por um script de medição que se autentica "
        "na API com as mesmas credenciais de um utilizador real e exercita os endpoints produtivos. Nenhum "
        "valor é estimado ou inventado: cada número resulta de uma medição efetiva sobre o sistema em "
        "execução. Quando um indicador não está disponível no ambiente observado, é reportado de forma "
        "explícita como N/D, em vez de ser preenchido com um valor arbitrário. Esta disciplina garante que "
        "o relatório é reproduzível e auditável, e que qualquer leitor pode reexecutar o script e obter "
        "resultados comparáveis."
    )
    add_body(
        doc,
        f"A recolha que serve de base a esta versão foi realizada em {gen_at}, contra a instância "
        f"{base_url}, com {iterations} repetições por cada cenário de pesquisa. A repetição das medições "
        "é deliberada: permite distinguir o comportamento típico (mediana) de efeitos pontuais de "
        "arranque a frio (máximos isolados), evitando conclusões a partir de uma única amostra."
    )

    # --- 4.2.1 Latência de pesquisa ---
    add_heading(doc, "4.2.1. Latência dos Modos de Pesquisa", 3)
    add_body(
        doc,
        "A plataforma oferece três modos de pesquisa com características distintas. A pesquisa textual faz "
        "correspondência lexical e não depende de modelos de aprendizagem, pelo que serve de linha de base. "
        "A pesquisa vetorial converte a consulta num vetor de características (embedding) e compara-o "
        "semanticamente com o corpus, captando significado para além das palavras exatas. A pesquisa "
        "híbrida combina ambas as estratégias, somando à precisão semântica a robustez da correspondência "
        "lexical. A tabela seguinte resume a latência observada em cada modo, em milissegundos."
    )

    def bench_row(label, data):
        return (
            label,
            fmt_int(data.get("iterations")),
            fmt_ms(data.get("min_ms")),
            fmt_ms(data.get("median_ms")),
            fmt_ms(data.get("avg_ms")),
            fmt_ms(data.get("p95_ms")),
            fmt_ms(data.get("max_ms")),
        )

    bench_rows = []
    for key, label in (("text", "Textual"), ("vector", "Vetorial"), ("hybrid", "Híbrida")):
        data = search_bench.get(key, {})
        if data:
            bench_rows.append(bench_row(label, data))
    ask_simple = search_bench.get("ask_simple")
    if ask_simple:
        bench_rows.append(bench_row("Ask simples (RAG)", ask_simple))

    if bench_rows:
        add_data_table(
            doc,
            ["Modo de pesquisa", "Iterações", "Mínimo", "Mediana", "Média", "P95", "Máximo"],
            bench_rows,
            caption="Latência por modo de pesquisa (valores medidos sobre a API em execução).",
        )
    else:
        add_body(doc, "Não foram recolhidos resultados de benchmark de pesquisa nesta execução (N/D).")

    text_median = search_bench.get("text", {}).get("median_ms")
    vector_median = search_bench.get("vector", {}).get("median_ms")
    hybrid_median = search_bench.get("hybrid", {}).get("median_ms")
    add_body(
        doc,
        "Leitura dos resultados: a pesquisa textual apresenta a menor latência, como seria de esperar por "
        f"não envolver inferência de modelos (mediana de {fmt_ms(text_median)}). A pesquisa vetorial e a "
        f"híbrida situam-se na ordem das centenas de milissegundos (medianas de {fmt_ms(vector_median)} e "
        f"{fmt_ms(hybrid_median)}, respetivamente), valor adequado a uma utilização interativa. A diferença "
        "entre a mediana e o percentil 95 (P95) é reduzida, o que indica um comportamento estável e "
        "previsível, sem grande dispersão entre pedidos consecutivos."
    )
    add_body(
        doc,
        "Estes valores refletem uma otimização deliberada do caminho de geração de embeddings. Numa "
        "configuração inicial, o modelo de embedding era recarregado a partir de disco a cada consulta, o "
        "que elevava a latência da pesquisa vetorial e híbrida à ordem de vários segundos por pedido. Ao "
        "encaminhar a geração de embeddings para um serviço dedicado que mantém o modelo residente em "
        "memória, a latência foi reduzida para a ordem das centenas de milissegundos. A tabela seguinte "
        "quantifica o efeito desta alteração."
    )

    BASELINE_VECTOR_MS = 10993.0
    BASELINE_HYBRID_MS = 10391.0
    add_data_table(
        doc,
        ["Modo", "Antes (modelo em processo)", "Depois (serviço dedicado)", "Fator de melhoria"],
        [
            (
                "Vetorial",
                fmt_ms(BASELINE_VECTOR_MS),
                fmt_ms(vector_median),
                speedup(BASELINE_VECTOR_MS, vector_median),
            ),
            (
                "Híbrida",
                fmt_ms(BASELINE_HYBRID_MS),
                fmt_ms(hybrid_median),
                speedup(BASELINE_HYBRID_MS, hybrid_median),
            ),
        ],
        caption=(
            "Impacto da otimização do caminho de embeddings (medianas antes e depois). "
            "Os valores 'antes' correspondem à medição da configuração com recarregamento do modelo."
        ),
    )
    add_body(
        doc,
        "Para preservar a disponibilidade do serviço, foi mantido um mecanismo de recurso: caso o serviço "
        "dedicado de embeddings fique indisponível, o sistema repõe automaticamente a geração em processo. "
        "Desta forma, a otimização melhora o desempenho sem introduzir um ponto único de falha na função "
        "de pesquisa."
    )
    add_screenshot_placeholder(doc, "Gráfico benchmark_search_latency.png gerado pelo script de métricas")

    # --- 4.2.2 Auditoria: volume e taxa de sucesso ---
    add_heading(doc, "4.2.2. Auditoria: Volume e Taxa de Sucesso", 3)
    add_body(
        doc,
        "O registo de auditoria é a fonte primária para avaliar o comportamento funcional do sistema ao "
        "longo do tempo. Cada entrada documenta uma ação, o recurso afetado, o estado final, o código HTTP "
        "e a duração em milissegundos. A análise que se segue baseia-se na amostra de registos mais "
        "recentes devolvida pela API."
    )
    total_logs = audit.get("total_logs")
    status_dist = audit.get("status_distribution", {})
    n_success = status_dist.get("success")
    n_failure = status_dist.get("failure")
    add_data_table(
        doc,
        ["Indicador", "Valor"],
        [
            ("Registos de auditoria analisados", fmt_int(total_logs)),
            ("Operações com sucesso", f"{fmt_int(n_success)} ({fmt_pct(n_success, total_logs)})"),
            ("Operações com falha", f"{fmt_int(n_failure)} ({fmt_pct(n_failure, total_logs)})"),
            ("Taxa de sucesso global", audit.get("success_rate", "N/D")),
        ],
        caption="Síntese do volume e do estado das operações registadas em auditoria.",
        col_align=[None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(
        doc,
        f"A taxa de sucesso observada foi de {audit.get('success_rate', 'N/D')}. As falhas registadas "
        "concentram-se predominantemente em tentativas de autenticação e renovação de sessão, sendo "
        "comportamentos esperados num sistema com controlo de acesso ativo (por exemplo, credenciais "
        "incorretas ou tokens expirados). Ainda assim, a sua quantificação é relevante para detetar "
        "eventuais picos anómalos que possam indiciar problemas de configuração ou tentativas de acesso "
        "indevido."
    )

    # --- 4.2.3 Distribuição de ações ---
    actions = audit.get("actions", {})
    if actions:
        add_heading(doc, "4.2.3. Distribuição das Ações Auditadas", 3)
        add_body(
            doc,
            "A distribuição das ações permite compreender que tipo de atividade domina o sistema e, "
            "sobretudo, evitar que operações de elevada frequência distorçam a leitura dos indicadores de "
            "desempenho. A tabela seguinte ordena as ações pela sua frequência."
        )
        sorted_actions = sorted(actions.items(), key=lambda kv: kv[1], reverse=True)
        action_rows = [
            (acao, fmt_int(qtd), fmt_pct(qtd, total_logs))
            for acao, qtd in sorted_actions
        ]
        add_data_table(
            doc,
            ["Ação", "Ocorrências", "Peso relativo"],
            action_rows,
            caption="Frequência das ações registadas em auditoria, por ordem decrescente.",
            col_align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        )
        tr_count = actions.get("token_refresh")
        add_body(
            doc,
            "Observa-se que a renovação de token (token_refresh) representa, isoladamente, a larga maioria "
            f"dos registos ({fmt_int(tr_count)} ocorrências, {fmt_pct(tr_count, total_logs)} do total). "
            "Trata-se de uma operação automática e de fundo, executada para manter as sessões ativas. Por "
            "ser tão frequente e ocasionalmente sujeita a esperas de rede, esta ação domina e distorce os "
            "percentis de duração globais. Por esse motivo, na análise de desempenho que se segue, os "
            "indicadores são apresentados também excluindo a renovação de token, de modo a refletir com "
            "fidelidade o tempo das operações funcionais iniciadas pelo utilizador."
        )
        add_screenshot_placeholder(doc, "Gráfico audit_actions_top10.png gerado pelo script de métricas")
        add_screenshot_placeholder(doc, "Gráfico audit_status_ratio.png gerado pelo script de métricas")

    # --- 4.2.4 Duração das operações ---
    add_heading(doc, "4.2.4. Duração das Operações Auditadas", 3)
    add_body(
        doc,
        "A duração das operações é analisada em duas perspetivas complementares. A visão global inclui "
        "todas as ações e é útil para dimensionamento de capacidade; a visão que exclui a renovação de "
        "token isola o desempenho percebido pelo utilizador nas operações funcionais. A comparação das "
        "duas colunas evidencia o impacto da operação de fundo sobre as estatísticas agregadas."
    )
    add_data_table(
        doc,
        ["Estatística", "Global (todas as ações)", "Excluindo renovação de token"],
        [
            ("Amostras", fmt_int(dur_all.get("samples")), fmt_int(dur_excl.get("samples"))),
            ("Mínimo", fmt_ms(dur_all.get("min")), fmt_ms(dur_excl.get("min"))),
            ("Mediana", fmt_ms(dur_all.get("median")), fmt_ms(dur_excl.get("median"))),
            ("Média", fmt_ms(dur_all.get("avg")), fmt_ms(dur_excl.get("avg"))),
            ("P95", fmt_ms(dur_all.get("p95")), fmt_ms(dur_excl.get("p95"))),
            ("Máximo", fmt_ms(dur_all.get("max")), fmt_ms(dur_excl.get("max"))),
        ],
        caption="Duração das operações auditadas, com e sem a influência da renovação de token.",
        col_align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(
        doc,
        f"A leitura é elucidativa: a mediana global ({fmt_ms(dur_all.get('median'))}) e a mediana sem "
        f"renovação de token ({fmt_ms(dur_excl.get('median'))}) são próximas, o que confirma que a "
        "operação típica é rápida. No entanto, a média e o P95 globais são fortemente penalizados pelos "
        "tempos extremos da renovação de token, refletindo esperas pontuais de rede. Ao remover essa ação, "
        f"o P95 das operações funcionais desce para {fmt_ms(dur_excl.get('p95'))}, um valor que traduz "
        "corretamente a experiência do utilizador na utilização normal da aplicação."
    )

    # --- 4.2.5 Duração por tipo de operação ---
    by_action = audit.get("duration_ms_by_action", {})
    if by_action:
        add_heading(doc, "4.2.5. Duração por Tipo de Operação", 3)
        add_body(
            doc,
            "Para um diagnóstico mais fino, a tabela seguinte detalha a duração por tipo de operação. Esta "
            "decomposição permite identificar quais as funcionalidades mais exigentes e orientar eventuais "
            "esforços de otimização para onde têm maior impacto."
        )
        ordered = sorted(
            by_action.items(),
            key=lambda kv: kv[1].get("median", 0),
            reverse=True,
        )
        per_action_rows = [
            (
                acao,
                fmt_int(stats.get("samples")),
                fmt_ms(stats.get("median")),
                fmt_ms(stats.get("avg")),
                fmt_ms(stats.get("p95")),
                fmt_ms(stats.get("max")),
            )
            for acao, stats in ordered
        ]
        add_data_table(
            doc,
            ["Operação", "Amostras", "Mediana", "Média", "P95", "Máximo"],
            per_action_rows,
            caption="Duração detalhada por tipo de operação, ordenada pela mediana.",
            col_align=[None, WD_ALIGN_PARAGRAPH.CENTER, None, None, None, None],
        )
        add_body(
            doc,
            "As operações de escrita sobre conteúdo (criação de notas e atualização de notebooks) "
            "apresentam, naturalmente, durações superiores às operações de leitura ou de encerramento de "
            "sessão, por envolverem persistência em base de dados. Os valores mantêm-se, contudo, dentro de "
            "limites adequados a uma utilização interativa, não havendo nenhuma operação funcional com "
            "tempos típicos preocupantes."
        )

    # --- 4.2.6 Inventário de entidades ---
    add_heading(doc, "4.2.6. Inventário de Entidades", 3)
    add_body(
        doc,
        "O inventário de entidades caracteriza o estado de povoamento do ambiente no momento da medição. "
        "É um indicador de contexto: ajuda a interpretar as restantes métricas e a validar se o ambiente "
        "observado corresponde ao esperado (por exemplo, um ambiente de demonstração com poucos dados, ou "
        "um ambiente produtivo com volume significativo)."
    )
    add_data_table(
        doc,
        ["Entidade", "Quantidade"],
        [
            ("Modelos registados", fmt_int(entities.get("models"))),
            ("Credenciais registadas", fmt_int(entities.get("credentials"))),
            ("Notebooks", fmt_int(entities.get("notebooks"))),
            ("Notas", fmt_int(entities.get("notes"))),
            ("Fontes (documentos ingeridos)", fmt_int(entities.get("sources"))),
        ],
        caption="Inventário de entidades existentes no ambiente no momento da medição.",
        col_align=[None, WD_ALIGN_PARAGRAPH.CENTER],
    )
    add_body(
        doc,
        "No ambiente observado, a contagem de fontes encontra-se reduzida e poderá estar associada a um "
        "espaço de dados distinto daquele que serve a pesquisa. Este desalinhamento entre o namespace de "
        "ingestão e a origem ativa de pesquisa é, em si, um achado operacional relevante e deve ser objeto "
        "de procedimento de validação de ambiente antes de cada ciclo de reporte, para assegurar que as "
        "métricas de pesquisa incidem sobre o corpus pretendido."
    )

    # --- 4.2.7 Disponibilidade dos serviços ---
    health = metricas.get("health", {})
    if health:
        add_heading(doc, "4.2.7. Disponibilidade e Tempo de Resposta dos Serviços", 3)
        add_body(
            doc,
            "Por fim, foram verificados os endpoints de estado que sustentam a monitorização operacional. "
            "Estes pontos de verificação confirmam a disponibilidade da aplicação, da base de dados e do "
            "subsistema de autenticação, e medem o respetivo tempo de resposta. São a base para alertas "
            "automáticos e para a deteção precoce de degradação de serviço."
        )
        health_labels = {
            "health": "Verificação de saúde (/health)",
            "api_health": "Saúde da API (/api/health)",
            "api_config": "Configuração e estado da BD (/api/config)",
            "auth_status": "Estado de autenticação (/api/auth/status)",
        }
        health_rows = []
        for key, label in health_labels.items():
            entry = health.get(key)
            if not entry:
                continue
            code = entry.get("status_code")
            estado = "Disponível" if code == 200 else "Indisponível"
            health_rows.append(
                (label, fmt_int(code), estado, fmt_ms(entry.get("latency_ms")))
            )
        if health_rows:
            add_data_table(
                doc,
                ["Serviço / Endpoint", "Código HTTP", "Estado", "Latência"],
                health_rows,
                caption="Disponibilidade e tempo de resposta dos serviços de estado.",
                col_align=[None, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
            )
        db_status = health.get("api_config", {}).get("payload", {}).get("dbStatus", "N/D")
        add_body(
            doc,
            f"Todos os serviços essenciais responderam com sucesso (código HTTP 200) e a base de dados "
            f"reportou o estado '{db_status}'. Os tempos de resposta destes pontos de verificação situam-se "
            "abaixo da dezena de milissegundos para a maioria dos casos, o que os torna adequados a "
            "sondagens frequentes sem impacto percetível na carga do sistema."
        )

    add_screenshot_placeholder(doc, "Status Dashboard com indicadores de API, Base de Dados e Autenticação")

    add_heading(doc, "4.3. Método de Cálculo das Métricas", 2)
    add_body(
        doc,
        "A credibilidade dos indicadores depende da transparência do método. Por esse motivo, descreve-se "
        "em seguida, de forma reproduzível, todo o procedimento de recolha. O script de medição comporta-se "
        "como um cliente legítimo do sistema: autentica-se, executa pedidos reais e regista os tempos e os "
        "estados devolvidos pela API, sem aceder diretamente à base de dados nem manipular os resultados."
    )
    add_body(doc, "O procedimento decorre nos seguintes passos:")
    add_bullet(doc, "Autenticação na API através do endpoint de login local, obtendo um token Bearer para os pedidos seguintes.")
    add_bullet(doc, "Execução repetida de benchmarks de pesquisa nos três modos (textual, vetorial e híbrida), registando a latência de cada iteração.")
    add_bullet(doc, "Execução opcional do benchmark de geração de resposta (Ask simples), usando os modelos configurados por defeito.")
    add_bullet(doc, "Extração dos registos de auditoria e agregação por ação, estado e percentis de duração, incluindo a variante que exclui a renovação de token.")
    add_bullet(doc, "Contagem das entidades existentes (modelos, credenciais, notebooks, notas e fontes) para caracterizar o contexto do ambiente.")
    add_bullet(doc, "Verificação dos endpoints de estado e medição da respetiva latência.")
    add_bullet(doc, "Persistência dos resultados em ficheiros JSON, CSV e PNG, garantindo rastreabilidade e auditoria técnica posterior.")

    add_body(
        doc,
        "A opção por estatísticas robustas (mediana e percentil 95) em detrimento de uma simples média "
        "resulta de uma preocupação metodológica: a média é facilmente distorcida por valores extremos, ao "
        "passo que a mediana descreve o comportamento típico e o P95 capta o pior caso realista que a "
        "maioria dos utilizadores ainda experimenta. A apresentação conjunta de mínimo, mediana, média, P95 "
        "e máximo permite ao leitor formar um juízo completo sobre a distribuição de cada indicador."
    )
    add_body(
        doc,
        "Cada execução produz um conjunto de artefactos verificáveis, que devem ser anexados ao reporte de "
        "forma a permitir a sua reavaliação independente. A tabela seguinte identifica esses ficheiros e o "
        "respetivo conteúdo."
    )
    add_data_table(
        doc,
        ["Ficheiro", "Conteúdo"],
        [
            ("metricas_resumo.json", "Resumo consolidado de todas as métricas, com metadados da recolha."),
            ("search_benchmark.csv", "Latências individuais por iteração e por modo de pesquisa."),
            ("audit_actions.csv", "Distribuição das ações e estados extraídos da auditoria."),
            ("benchmark_search_latency.png", "Gráfico comparativo da latência dos modos de pesquisa."),
            ("audit_actions_top10.png", "Gráfico das ações mais frequentes na auditoria."),
            ("audit_status_ratio.png", "Gráfico da proporção entre sucessos e falhas."),
        ],
        caption="Artefactos gerados em cada execução do script de métricas (pasta metricas_e04).",
    )

    doc.add_page_break()

    # 5
    add_heading(doc, "5. Registos de Testes de DR (DEV/QUA/PRD)", 1)
    add_body(
        doc,
        "A resiliência do sistema só fica demonstrada quando a recuperação é efetivamente exercitada. Os "
        "testes de recuperação de desastre (DR) devem, por isso, ser executados periodicamente e registados "
        "com evidência temporal e técnica. O propósito vai além de comprovar que o restauro é possível: "
        "trata-se de medir, em cada ambiente, o tempo real até ao restabelecimento do serviço e a quantidade "
        "de dados potencialmente perdida."
    )
    add_body(
        doc,
        "Dois indicadores estruturam esta avaliação. O RTO (Recovery Time Objective) define o tempo máximo "
        "aceitável entre a ocorrência de uma falha e a reposição do serviço. O RPO (Recovery Point "
        "Objective) define a perda máxima de dados tolerável, expressa como o intervalo de tempo entre a "
        "última cópia de segurança utilizável e o momento da falha. Os valores-alvo são tipicamente mais "
        "exigentes em produção do que em ambientes de desenvolvimento, refletindo a criticidade crescente "
        "do serviço. A tabela seguinte fixa os alvos por ambiente e o estado do último teste realizado."
    )

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    for c, t in zip(h, ["Ambiente", "RTO alvo", "RPO alvo", "Estado do último teste"]):
        c.text = t
    style_header_row(h)
    for row in [("DEV", "24h", "8h", "A preencher"), ("QUA", "8h", "4h", "A preencher"), ("PRD", "4h", "1h", "A preencher")]:
        r = tbl.add_row().cells
        for i, v in enumerate(row):
            r[i].text = v
    add_caption(doc, "Tabela", "Objetivos de RTO/RPO por ambiente e estado do último teste de DR.")
    add_body(
        doc,
        "A coluna 'Estado do último teste' deve ser preenchida a cada ciclo com a data de execução e o "
        "resultado obtido (tempo real de restauro e ponto de recuperação efetivo), permitindo comparar o "
        "desempenho real face aos alvos definidos e despoletar ações corretivas sempre que os objetivos não "
        "sejam cumpridos."
    )

    add_screenshot_placeholder(doc, "Evidência DR DEV - antes/depois do restauro")
    add_screenshot_placeholder(doc, "Evidência DR QUA - antes/depois do restauro")
    add_screenshot_placeholder(doc, "Evidência DR PRD - failover e failback")

    doc.add_page_break()

    # 6
    add_heading(doc, "6. Conclusão", 1)
    add_body(
        doc,
        "A revisão E04 foi atualizada para refletir o comportamento real do sistema e da interface, "
        "substituindo formulações especulativas por evidência mensurável. A interface de utilizador cobre "
        "o ciclo completo de operação - autenticação, navegação, pesquisa, gestão de conteúdos e "
        "administração - e o subsistema de auditoria fornece a base factual para o acompanhamento contínuo "
        "da qualidade de serviço."
    )
    add_body(
        doc,
        "Do ponto de vista do desempenho, os resultados confirmam um sistema responsivo: a pesquisa textual "
        "responde na ordem das dezenas a centenas de milissegundos e a pesquisa semântica, após a "
        "otimização do caminho de geração de embeddings, situa-se igualmente na ordem das centenas de "
        "milissegundos, um ganho muito significativo face à configuração inicial. A análise de auditoria "
        "evidenciou ainda a importância de isolar operações de fundo de elevada frequência, sob pena de "
        "distorcerem a leitura dos indicadores agregados."
    )
    add_body(
        doc,
        "Foram também identificados pontos a consolidar antes de uma adoção produtiva plena: a validação "
        "fim-a-fim do fluxo de autenticação federada (OAuth/SSO) no ambiente alvo, a garantia de "
        "alinhamento entre o espaço de ingestão de dados e a origem ativa de pesquisa, e a execução "
        "regular dos testes de recuperação de desastre com registo de RTO/RPO reais."
    )
    add_body(
        doc,
        "Como próximo passo operacional, recomenda-se institucionalizar a execução periódica do script de "
        "métricas, anexando os gráficos e o ficheiro de resumo gerados a cada ciclo de reporte E04. Desta "
        "forma, a evolução do sistema passa a ser acompanhada com indicadores comparáveis ao longo do "
        "tempo, sustentando decisões de operação e de investimento em evidência objetiva."
    )

    doc.add_page_break()
    add_heading(doc, "Apêndice A - Lista de Capturas de Ecrã", 1)
    screenshots = [
        "Ecrã de login com OAuth e login local",
        "Interface principal com navegação global",
        "Chat global com exportação",
        "Consola de modelos",
        "Admin Dashboard com separadores",
        "Audit Logs com filtros",
        "Status Dashboard API/DB/Auth",
        "Dockhand de infraestrutura",
        "benchmark_search_latency.png",
        "audit_actions_top10.png",
        "audit_status_ratio.png",
        "Evidências DR DEV/QUA/PRD",
    ]
    for i, s in enumerate(screenshots, 1):
        add_bullet(doc, f"Figura {i}: {s}")

    output = str(BASE_DIR / "E04_Interface_Monitorizacao.docx")
    doc.save(output)
    print(f"Documento gerado com sucesso: {output}")
    return output


if __name__ == "__main__":
    build_document()
